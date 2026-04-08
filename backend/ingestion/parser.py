"""
Document parser — converts PDF / DOCX / TXT files into clean, structured text.

PDF strategy (dual-pass):
  Pass 1 — pdfplumber detects tables and converts them to Label: Value lines.
  Pass 2 — non-table text is extracted via pdfplumber; PyMuPDF used as fallback
            when pdfplumber yields nothing for a page.

This preserves relationships like:
  Shipper: AAA, LAX...
  Consignee: xyz, Fontana...
instead of the scrambled multi-column output from naive left-to-right extraction.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF fallback
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


# ── Public entry point ─────────────────────────────────────────────────────────

def parse_document(filepath: str | Path) -> dict:
    """
    Parse a document and return:
      {
        "raw_text": str,      # full structured text ready for chunking
        "page_count": int,
        "doc_type": str,      # "pdf" | "docx" | "txt"
      }
    """
    path = Path(filepath)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        text, pages = _parse_pdf(path)
        doc_type = "pdf"
    elif suffix == ".docx":
        text, pages = _parse_docx(path)
        doc_type = "docx"
    elif suffix == ".txt":
        text, pages = _parse_txt(path)
        doc_type = "txt"
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    text = text.strip()
    if len(text) < 30:
        raise ValueError(
            "Document appears to be image-based or empty. "
            "Text extraction yielded insufficient content."
        )

    return {"raw_text": text, "page_count": pages, "doc_type": doc_type}


# ── PDF ────────────────────────────────────────────────────────────────────────

def _parse_pdf(path: Path) -> tuple[str, int]:
    blocks: list[str] = []
    page_count = 0

    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)

        for page in pdf.pages:
            page_blocks: list[str] = []

            # --- Pass 1: extract tables as key-value text ---
            try:
                table_finder = page.find_tables()
                table_bboxes = [t.bbox for t in table_finder]
                raw_tables   = page.extract_tables()

                for raw_table in raw_tables:
                    kv = _table_to_kv(raw_table)
                    if kv:
                        page_blocks.append("\n".join(kv))
            except Exception as exc:
                logger.warning("Table extraction failed on page %s: %s", page.page_number, exc)
                table_bboxes = []

            # --- Pass 2: extract non-table text ---
            try:
                # Crop away table regions and extract the remaining text
                remaining_page = page
                for bbox in table_bboxes:
                    try:
                        remaining_page = remaining_page.outside_bbox(bbox)
                    except Exception:
                        pass
                plain = remaining_page.extract_text(x_tolerance=3, y_tolerance=3)
                if plain and plain.strip():
                    page_blocks.append(plain.strip())
            except Exception as exc:
                logger.warning("Text extraction failed on page %s: %s", page.page_number, exc)

            # --- PyMuPDF fallback if pdfplumber got nothing for this page ---
            if not page_blocks:
                fallback = _pymupdf_page_text(path, page.page_number - 1)
                if fallback:
                    page_blocks.append(fallback)

            if page_blocks:
                blocks.append("\n\n".join(page_blocks))

    return "\n\n---\n\n".join(blocks), page_count


def _pymupdf_page_text(path: Path, page_index: int) -> Optional[str]:
    try:
        doc = fitz.open(str(path))
        page = doc.load_page(page_index)
        text = page.get_text("text")
        doc.close()
        return text.strip() or None
    except Exception as exc:
        logger.warning("PyMuPDF fallback failed: %s", exc)
        return None


# ── Table → key-value conversion ───────────────────────────────────────────────

def _table_to_kv(table: list[list]) -> list[str]:
    """
    Convert a 2-D table (list of rows, each row a list of cell strings) into
    a flat list of "Label: Value" strings.

    Handles three layouts:
      A) Vertical key-value (2-col, each row is Label | Value)
         e.g. Load ID | LD53657 / Ship Date | 02-08-2026 09:00
      B) Horizontal header + data rows (>2 cols or 2-col where col2 of row1 looks like a label)
         e.g. # Of Units | Description | Weight | Type | Class  (header row)
      C) Multi-column data row → "Col1 | Col2 | Col3"
    """
    if not table:
        return []

    # Normalise cells: None → ""
    table = [[_clean(c) for c in row] for row in table]

    # Remove fully-empty rows
    table = [row for row in table if any(c for c in row)]
    if not table:
        return []

    lines: list[str] = []
    first_row = table[0]
    non_empty_first = [c for c in first_row if c]
    num_cols = max(len(row) for row in table)

    # ── Determine layout ───────────────────────────────────────────────────────
    #
    # For 2-column tables, distinguish:
    #   Vertical KV:  row1 = ["Load ID", "LD53657"]  → col2 looks like a value
    #   Horizontal:   row1 = ["Shipper", "Consignee"] → col2 looks like a label
    #
    # For >2-column tables, row1 is always treated as a header.

    if num_cols == 2 and len(table) > 1:
        # If column 2 of the first row looks like a data value → vertical KV layout
        col2_first = first_row[1] if len(first_row) > 1 else ""
        if _looks_like_value(col2_first):
            # Each row is: Label | Value
            for row in table:
                non_empty = [c for c in row if c]
                if len(non_empty) == 2:
                    lines.append(f"{non_empty[0]}: {non_empty[1]}")
                elif len(non_empty) == 1:
                    lines.append(non_empty[0])
            return lines
        else:
            # First row is a header (e.g., Shipper | Consignee)
            headers = first_row
            for row in table[1:]:
                for i, cell in enumerate(row):
                    if cell and i < len(headers) and headers[i]:
                        lines.append(f"{headers[i]}: {cell}")
            return lines

    # >2 columns: first row is a header
    has_header = (
        len(table) > 1
        and all(not _is_numeric(c) for c in non_empty_first if c)
        and len(non_empty_first) > 0
    )

    if has_header:
        headers = first_row
        for row in table[1:]:
            for i, cell in enumerate(row):
                if cell and i < len(headers) and headers[i]:
                    lines.append(f"{headers[i]}: {cell}")
    else:
        for row in table:
            non_empty = [c for c in row if c]
            if len(non_empty) == 0:
                continue
            elif len(non_empty) == 1:
                lines.append(non_empty[0])
            elif len(non_empty) == 2:
                lines.append(f"{non_empty[0]}: {non_empty[1]}")
            else:
                lines.append(" | ".join(non_empty))

    return lines


def _looks_like_value(s: str) -> bool:
    """
    Returns True if a string looks like a data value rather than a column label.
    Used to distinguish vertical KV tables from horizontal header tables.
    """
    if not s:
        return False
    # Contains digits (dates, IDs, amounts) → likely a value
    if any(ch.isdigit() for ch in s):
        return True
    # Long text (addresses, descriptions) → likely a value
    if len(s) > 25:
        return True
    return False


def _clean(cell) -> str:
    if cell is None:
        return ""
    return str(cell).strip()


def _is_numeric(s: str) -> bool:
    return s.lstrip("#$-+").replace(".", "").replace(",", "").isnumeric()


# ── DOCX ───────────────────────────────────────────────────────────────────────

def _parse_docx(path: Path) -> tuple[str, int]:
    doc = Document(str(path))
    blocks: list[str] = []
    table_index = 0

    # Walk body elements in document order to preserve layout
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            text = element.text_content() if hasattr(element, "text_content") else ""
            # python-docx: use the Paragraph wrapper for clean text
            try:
                from docx.oxml.ns import qn
                runs = element.findall(f".//{qn('w:t')}")
                text = " ".join(r.text for r in runs if r.text)
            except Exception:
                pass
            if text.strip():
                blocks.append(text.strip())

        elif tag == "tbl":
            if table_index < len(doc.tables):
                kv = _docx_table_to_kv(doc.tables[table_index])
                if kv:
                    blocks.append("\n".join(kv))
                table_index += 1

    return "\n\n".join(blocks), 1


def _docx_table_to_kv(table) -> list[str]:
    rows = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        rows.append(cells)
    return _table_to_kv(rows)


# ── TXT ────────────────────────────────────────────────────────────────────────

def _parse_txt(path: Path) -> tuple[str, int]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return text, 1
